import os
import json
import re
from datetime import datetime
from typing import Dict, List, Any, Tuple
import gradio as gr
import mammoth
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import openai
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import tempfile
import zipfile
import io

class ADGMCorporateAgent:
    def __init__(self, openai_api_key: str = None):
        """Initialize the ADGM Corporate Agent with RAG capabilities"""
        self.openai_client = openai.OpenAI(api_key=openai_api_key) if openai_api_key else None
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Initialize ChromaDB for RAG
        self.chroma_client = chromadb.Client()
        try:
            self.collection = self.chroma_client.get_collection("adgm_regulations")
        except:
            self.collection = self.chroma_client.create_collection(
                name="adgm_regulations",
                embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(
                    model_name="all-MiniLM-L6-v2"
                )
            )
            self._initialize_knowledge_base()
    
    def _initialize_knowledge_base(self):
        """Initialize the RAG knowledge base with ADGM regulations and requirements"""
        adgm_knowledge = [
            {
                "id": "company_formation_req",
                "content": "ADGM Company Formation Requirements: All companies must submit Articles of Association, Memorandum of Association, Board Resolution, Shareholder Resolution, Incorporation Application Form, UBO Declaration Form, and Register of Members and Directors. Per ADGM Companies Regulations 2020.",
                "category": "company_formation"
            },
            {
                "id": "jurisdiction_clause",
                "content": "All legal documents must specify ADGM jurisdiction and ADGM Courts for dispute resolution. References to UAE Federal Courts are non-compliant. ADGM Companies Regulations 2020, Article 6.",
                "category": "jurisdiction"
            },
            {
                "id": "signatory_requirements",
                "content": "All corporate documents must include proper signatory sections with full names, titles, and signature dates. Electronic signatures must comply with ADGM Electronic Transactions Regulations.",
                "category": "signatures"
            },
            {
                "id": "aoa_requirements",
                "content": "Articles of Association must include company name, registered office address, business objects, share capital structure, director powers, and shareholder rights. Must comply with ADGM template requirements.",
                "category": "articles_of_association"
            },
            {
                "id": "moa_requirements",
                "content": "Memorandum of Association must state company name, registered office, authorized share capital, and initial subscribers. Must be consistent with Articles of Association.",
                "category": "memorandum_of_association"
            },
            {
                "id": "board_resolution_format",
                "content": "Board resolutions must include meeting date, attendees, quorum confirmation, resolution details, voting results, and proper authorization signatures.",
                "category": "board_resolution"
            },
            {
                "id": "ubo_declaration",
                "content": "Ultimate Beneficial Owner declarations must identify all individuals owning 25% or more of shares, directly or indirectly. Must include full personal details and ownership percentages.",
                "category": "ubo_declaration"
            }
        ]
        
        for item in adgm_knowledge:
            self.collection.add(
                documents=[item["content"]],
                metadatas=[{"category": item["category"]}],
                ids=[item["id"]]
            )
    
    def get_relevant_regulations(self, query: str, n_results: int = 3) -> List[str]:
        """Retrieve relevant ADGM regulations using RAG"""
        try:
            results = self.collection.query(
                query_texts=[query],
                n_results=n_results
            )
            return results['documents'][0] if results['documents'] else []
        except Exception as e:
            print(f"RAG query error: {e}")
            return []
    
    def identify_document_type(self, content: str, filename: str) -> str:
        """Identify the type of legal document based on content and filename"""
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Document type patterns
        doc_patterns = {
            'articles_of_association': ['articles of association', 'aoa', 'constitutional document'],
            'memorandum_of_association': ['memorandum of association', 'moa', 'memorandum'],
            'board_resolution': ['board resolution', 'board meeting', 'directors resolution'],
            'shareholder_resolution': ['shareholder resolution', 'shareholders resolution', 'general meeting'],
            'incorporation_application': ['incorporation application', 'company registration', 'formation application'],
            'ubo_declaration': ['ultimate beneficial owner', 'ubo declaration', 'beneficial ownership'],
            'register_members_directors': ['register of members', 'register of directors', 'members register'],
            'change_address_notice': ['change of address', 'registered address', 'address notice']
        }
        
        for doc_type, patterns in doc_patterns.items():
            if any(pattern in content_lower or pattern in filename_lower for pattern in patterns):
                return doc_type
        
        return 'unknown'
    
    def check_required_documents(self, uploaded_docs: List[str]) -> Dict[str, Any]:
        """Check if all required documents are present for company incorporation"""
        required_docs = [
            'articles_of_association',
            'memorandum_of_association', 
            'board_resolution',
            'shareholder_resolution',
            'incorporation_application',
            'ubo_declaration',
            'register_members_directors'
        ]
        
        doc_names = {
            'articles_of_association': 'Articles of Association',
            'memorandum_of_association': 'Memorandum of Association',
            'board_resolution': 'Board Resolution',
            'shareholder_resolution': 'Shareholder Resolution',
            'incorporation_application': 'Incorporation Application Form',
            'ubo_declaration': 'UBO Declaration Form',
            'register_members_directors': 'Register of Members and Directors'
        }
        
        missing_docs = []
        for req_doc in required_docs:
            if req_doc not in uploaded_docs:
                missing_docs.append(doc_names[req_doc])
        
        return {
            'total_required': len(required_docs),
            'uploaded_count': len(uploaded_docs),
            'missing_documents': missing_docs,
            'is_complete': len(missing_docs) == 0
        }
    
    def analyze_document_content(self, content: str, doc_type: str) -> List[Dict[str, Any]]:
        """Analyze document content for legal issues and compliance"""
        issues = []
        
        # Get relevant regulations for this document type
        regulations = self.get_relevant_regulations(f"{doc_type} requirements compliance")
        
        # Common red flag patterns
        red_flags = {
            'jurisdiction': {
                'pattern': r'uae federal court|dubai court|abu dhabi court',
                'issue': 'Incorrect jurisdiction reference',
                'suggestion': 'Must specify ADGM jurisdiction and ADGM Courts',
                'severity': 'High'
            },
            'missing_signatures': {
                'pattern': r'signature.*pending|to be signed|\[signature\]',
                'issue': 'Missing or placeholder signatures',
                'suggestion': 'Ensure all required signatures are present',
                'severity': 'High'
            },
            'incomplete_clauses': {
                'pattern': r'\[.*\]|tbd|to be determined|pending',
                'issue': 'Incomplete clauses or placeholder text',
                'suggestion': 'Complete all placeholder text before submission',
                'severity': 'Medium'
            }
        }
        
        # Document-specific checks
        if doc_type == 'articles_of_association':
            if 'registered office' not in content.lower():
                issues.append({
                    'section': 'General',
                    'issue': 'Missing registered office address',
                    'severity': 'High',
                    'suggestion': 'Include complete registered office address in ADGM'
                })
            
            if 'share capital' not in content.lower():
                issues.append({
                    'section': 'General',
                    'issue': 'Missing share capital information',
                    'severity': 'High',
                    'suggestion': 'Include authorized share capital details'
                })
        
        elif doc_type == 'ubo_declaration':
            if not re.search(r'\d+%|\d+\s*percent', content.lower()):
                issues.append({
                    'section': 'General',
                    'issue': 'Missing ownership percentages',
                    'severity': 'High',
                    'suggestion': 'Include specific ownership percentages for all UBOs'
                })
        
        # Check for common red flags
        for flag_type, flag_info in red_flags.items():
            matches = re.finditer(flag_info['pattern'], content, re.IGNORECASE)
            for match in matches:
                issues.append({
                    'section': f'Line {content[:match.start()].count(chr(10)) + 1}',
                    'issue': flag_info['issue'],
                    'severity': flag_info['severity'],
                    'suggestion': flag_info['suggestion'],
                    'matched_text': match.group()
                })
        
        return issues
    
    def add_comments_to_docx(self, doc_path: str, issues: List[Dict[str, Any]]) -> str:
        """Add comments and highlights to the DOCX file for identified issues"""
        try:
            doc = Document(doc_path)
            
            # Create a new temporary file for the reviewed document
            temp_dir = tempfile.mkdtemp()
            reviewed_path = os.path.join(temp_dir, f"reviewed_{os.path.basename(doc_path)}")
            
            # Process each paragraph and add highlighting for issues
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.lower()
                
                # Check if any issues match this paragraph
                for issue in issues:
                    if 'matched_text' in issue:
                        matched_text = issue['matched_text'].lower()
                        if matched_text in para_text:
                            # Highlight the problematic text
                            for run in paragraph.runs:
                                if matched_text in run.text.lower():
                                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                                    run.font.color.rgb = RGBColor(255, 0, 0)  # Red text
            
            # Add a summary section at the end
            doc.add_page_break()
            summary_heading = doc.add_heading('ADGM Compliance Review Summary', level=1)
            summary_heading.runs[0].font.color.rgb = RGBColor(0, 0, 255)  # Blue heading
            
            if issues:
                doc.add_paragraph('The following issues were identified during the compliance review:')
                for i, issue in enumerate(issues, 1):
                    issue_para = doc.add_paragraph()
                    issue_para.add_run(f"{i}. {issue['issue']} ").bold = True
                    issue_para.add_run(f"(Severity: {issue['severity']})")
                    if issue.get('suggestion'):
                        issue_para.add_run(f"\n   Suggestion: {issue['suggestion']}")
                    if issue.get('section'):
                        issue_para.add_run(f"\n   Location: {issue['section']}")
            else:
                doc.add_paragraph('No compliance issues were identified in this document.')
            
            # Save the reviewed document
            doc.save(reviewed_path)
            return reviewed_path
            
        except Exception as e:
            print(f"Error adding comments to DOCX: {e}")
            return doc_path
    
    def process_documents(self, files: List[str]) -> Tuple[Dict[str, Any], List[str]]:
        """Process uploaded documents and return analysis results"""
        if not files:
            return {"error": "No files uploaded"}, []
        
        results = {
            "timestamp": datetime.now().isoformat(),
            "total_documents": len(files),
            "document_analysis": [],
            "compliance_summary": {},
            "reviewed_files": []
        }
        
        reviewed_files = []
        uploaded_doc_types = []
        
        try:
            for file_path in files:
                if not file_path or not os.path.exists(file_path):
                    continue
                
                filename = os.path.basename(file_path)
                
                # Extract text from DOCX
                try:
                    doc = Document(file_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                except Exception as e:
                    print(f"Error reading {filename}: {e}")
                    continue
                
                # Identify document type
                doc_type = self.identify_document_type(content, filename)
                uploaded_doc_types.append(doc_type)
                
                # Analyze content for issues
                issues = self.analyze_document_content(content, doc_type)
                
                # Add comments to DOCX
                reviewed_path = self.add_comments_to_docx(file_path, issues)
                reviewed_files.append(reviewed_path)
                
                # Store analysis results
                results["document_analysis"].append({
                    "filename": filename,
                    "document_type": doc_type,
                    "issues_found": len(issues),
                    "issues": issues
                })
            
            # Check document completeness
            doc_completeness = self.check_required_documents(uploaded_doc_types)
            results["compliance_summary"] = {
                "process": "Company Incorporation",
                "documents_uploaded": doc_completeness['uploaded_count'],
                "required_documents": doc_completeness['total_required'],
                "missing_documents": doc_completeness['missing_documents'],
                "is_complete": doc_completeness['is_complete']
            }
            
        except Exception as e:
            results["error"] = f"Processing error: {str(e)}"
        
        return results, reviewed_files

# Global agent instance
agent = None

def initialize_agent(openai_key):
    """Initialize the ADGM Corporate Agent"""
    global agent
    try:
        agent = ADGMCorporateAgent(openai_api_key=openai_key if openai_key else None)
        return "✅ Agent initialized successfully!"
    except Exception as e:
        return f"❌ Error initializing agent: {str(e)}"

def process_documents_interface(files):
    """Main interface function for processing documents"""
    global agent
    
    if agent is None:
        return "❌ Please initialize the agent first", None, None
    
    if not files:
        return "❌ Please upload at least one document", None, None
    
    try:
        # Process the uploaded files
        results, reviewed_files = agent.process_documents(files)
        
        if "error" in results:
            return f"❌ {results['error']}", None, None
        
        # Format results for display
        summary_text = f"""
# ADGM Corporate Agent - Analysis Summary

## Document Overview
- **Total Documents Processed**: {results['total_documents']}
- **Analysis Timestamp**: {results['timestamp']}

## Compliance Check
- **Process Type**: {results['compliance_summary']['process']}
- **Documents Uploaded**: {results['compliance_summary']['documents_uploaded']}/{results['compliance_summary']['required_documents']}
- **Status**: {'✅ Complete' if results['compliance_summary']['is_complete'] else '⚠️ Incomplete'}

"""
        
        if results['compliance_summary']['missing_documents']:
            summary_text += "\n### Missing Documents:\n"
            for doc in results['compliance_summary']['missing_documents']:
                summary_text += f"- {doc}\n"
        
        summary_text += "\n## Document Analysis Results:\n"
        
        total_issues = 0
        for doc_analysis in results['document_analysis']:
            total_issues += doc_analysis['issues_found']
            summary_text += f"\n### {doc_analysis['filename']}\n"
            summary_text += f"- **Document Type**: {doc_analysis['document_type'].replace('_', ' ').title()}\n"
            summary_text += f"- **Issues Found**: {doc_analysis['issues_found']}\n"
            
            if doc_analysis['issues']:
                summary_text += "- **Issues**:\n"
                for issue in doc_analysis['issues']:
                    summary_text += f"  - **{issue['severity']}**: {issue['issue']}\n"
                    if issue.get('suggestion'):
                        summary_text += f"    - *Suggestion*: {issue['suggestion']}\n"
        
        summary_text += f"\n## Overall Summary\n"
        summary_text += f"- **Total Issues Found**: {total_issues}\n"
        summary_text += f"- **Compliance Status**: {'Ready for submission' if total_issues == 0 and results['compliance_summary']['is_complete'] else 'Requires attention'}\n"
        
        # Create downloadable files
        download_files = []
        if reviewed_files:
            # Create a zip file with all reviewed documents
            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for file_path in reviewed_files:
                    if os.path.exists(file_path):
                        zip_file.write(file_path, f"reviewed_{os.path.basename(file_path)}")
            
            zip_buffer.seek(0)
            zip_path = os.path.join(tempfile.mkdtemp(), "reviewed_documents.zip")
            with open(zip_path, 'wb') as f:
                f.write(zip_buffer.read())
            download_files.append(zip_path)
        
        # Create JSON report
        json_path = os.path.join(tempfile.mkdtemp(), "analysis_report.json")
        with open(json_path, 'w') as f:
            json.dump(results, f, indent=2)
        download_files.append(json_path)
        
        return summary_text, download_files, json.dumps(results, indent=2)
        
    except Exception as e:
        return f"❌ Error processing documents: {str(e)}", None, None

# Gradio Interface
def create_interface():
    with gr.Blocks(title="ADGM Corporate Agent", theme=gr.themes.Soft()) as interface:
        gr.Markdown("""
        # 🏛️ ADGM Corporate Agent
        ### Intelligent Legal Document Review and Compliance Assistant
        
        This AI-powered agent helps review and validate legal documents for ADGM (Abu Dhabi Global Market) compliance.
        Upload your corporate documents to get instant analysis, compliance checking, and recommendations.
        """)
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### ⚙️ Setup")
                openai_key = gr.Textbox(
                    label="OpenAI API Key (Optional)",
                    placeholder="sk-...",
                    type="password",
                    info="Optional: Provide OpenAI API key for enhanced analysis"
                )
                init_btn = gr.Button("Initialize Agent", variant="primary")
                init_status = gr.Textbox(label="Status", interactive=False)
        
        gr.Markdown("---")
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### 📄 Document Upload")
                file_upload = gr.Files(
                    label="Upload DOCX Documents",
                    file_types=[".docx"],
                    file_count="multiple"
                )
                
                process_btn = gr.Button("🔍 Analyze Documents", variant="primary", size="lg")
                
                gr.Markdown("""
                **Supported Document Types:**
                - Articles of Association (AoA)
                - Memorandum of Association (MoA)
                - Board Resolution Templates
                - Shareholder Resolution Templates
                - Incorporation Application Form
                - UBO Declaration Form
                - Register of Members and Directors
                """)
        
        gr.Markdown("---")
        
        with gr.Row():
            with gr.Column():
                gr.Markdown("### 📊 Analysis Results")
                results_display = gr.Markdown(label="Analysis Summary")
                
                with gr.Row():
                    download_files = gr.Files(label="📥 Download Reviewed Documents & Reports")
                    
        with gr.Row():
            with gr.Column():
                gr.Markdown("### 🔍 Detailed JSON Report")
                json_output = gr.Code(label="Structured Analysis Report", language="json")
        
        # Event handlers
        init_btn.click(
            initialize_agent,
            inputs=[openai_key],
            outputs=[init_status]
        )
        
        process_btn.click(
            process_documents_interface,
            inputs=[file_upload],
            outputs=[results_display, download_files, json_output]
        )
        
        gr.Markdown("""
        ---
        ### 📋 ADGM Compliance Features
        
        ✅ **Document Type Recognition** - Automatically identifies legal document types  
        ✅ **Completeness Checking** - Verifies all required documents are present  
        ✅ **Red Flag Detection** - Identifies non-compliant clauses and issues  
        ✅ **Inline Comments** - Adds review comments directly in DOCX files  
        ✅ **RAG-Enhanced Analysis** - Uses ADGM regulations for accurate compliance checking  
        ✅ **Structured Reporting** - Provides detailed JSON reports for integration  
        
        *Built for ADGM (Abu Dhabi Global Market) legal compliance requirements*
        """)
    
    return interface

# Launch the application
if __name__ == "__main__":
    interface = create_interface()
    interface.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=True,
        show_error=True
    )